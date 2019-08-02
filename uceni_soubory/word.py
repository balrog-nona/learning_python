import docx

# podle Automate the boring stuff with Python - lekce 13

"""
data types:
1. Document object - represents the entire document. Contains a list of:
2. Paragraph objects - for the paragraphs; new paragraph begins whenever the user presses enter or return while typing
Each paragraph object contains a list of:
3. Run objects - contiguous run of text with the same style. A new object is needed whenever the text style changes
"""

# reading document
doc = docx.Document('../../../automate_online-materials/demo.docx')
print(len(doc.paragraphs))  # 7 paragraph objects in the document
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(len(doc.paragraphs[1].runs))  # pocet run objectu v prvnim paragraphu

for item in doc.paragraphs[1].runs:
    print(item.text)

print('---------------------------------------------')

# pro prostou extrakci textu bez ohledu na jeho cleneni si napisu fci getText()
def getText(finemane):
    """
    Da se prizpusobit i tak, aby text nejak modifikovala
    Samozrejme se s extrahovanymtextem da delat cokoli - vyhledavat apod.
    """
    doc = docx.Document(finemane)
    fulltext = []
    for para in doc.paragraphs:
        """
        iteraci pres paragraphs se da extrahovat text
        """
        fulltext.append(para.text)
    return '\n'.join(fulltext)

print(getText('../../../automate_online-materials/demo.docx'))

# styling
"""
Paragraph styles can be applied to Paragraph objects
Character styles can be applied to Run objects
linked styles to both
dela se to tak, ze priradim atribut style to Paragraph or Run object. je preddefinovane, jak ten atribut muze znit

Runs can be further styled using text attributes - 3 values: True, False, None
"""
doc = docx.Document('../../../automate_online-materials/demo.docx')
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)
print(doc.paragraphs[1].text)

for item in doc.paragraphs[1].runs:
    # dostanu se k tomu, jak je odstavec rozdeleny na casti, se kteryma se da manipulovat
    print(item.text)

doc.paragraphs[0].style = 'Normal'  # zmena stylu z nadpisu na normalni text
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('../../../automate_online-materials/restyled.docx')  # vytvoril se novy dokument


# tvorba dokumentu
doc = docx.Document()
doc.add_paragraph('Hello world!')
doc.save('../../../automate_online-materials/helloworld.docx')

doc = docx.Document()
doc.add_paragraph('Hello world!')
paraObj1 = doc.add_paragraph('This is a second paragraph.')  # pridavat odstavec lze pouze na konec dokumentu
paraObj2 = doc.add_paragraph('Yet another paragraph.')
# pridavat run object lze pouze na konec paragraph objectu
paraObj1.add_run('This text is being added to the second paragraph.')
doc.save('../../../automate_online-materials/multipleparagraphs.docx')
# add_paragraph() a add_run() muzou mit i druhy argument - string s urcenim stylu

# adding headings
doc = docx.Document()
doc.add_heading('Header 0', 0)  # int urcuje title style, resp. uroven nadpisu v dokumentu
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('../../../automate_online-materials/headings.docx')


# lze vkladat obrazky na konec souboru pomoci add_picture()


